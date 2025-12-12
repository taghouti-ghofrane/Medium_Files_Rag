# -*- coding: utf-8 -*-
# md -> pdf
# md -> docx
import os
import sys
import argparse
from typing import List, Tuple
import markdown
# For DOCX
from docx import Document
from docx.shared import Inches
# For PDF
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table as RLTable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors
from reportlab.platypus import TableStyle
from PIL import Image as PILImage
import subprocess
import os
import sys
from typing import Optional
import re # For simple image path extraction from MD

# --- 1. Markdown parsing and content extraction ---
def extract_content_from_md(md_text: str) -> Tuple[List[str], List[str], List[List[List[str]]]]:
    """
    Extract plain text paragraphs, image paths, and table data from Markdown text.
    Note: This is a simplified parser suitable for basic MD format.
    For complex MD (like nested lists, code blocks), may need more powerful libraries like `markdown` + HTML parsing.
    """
    lines = md_text.strip().split('\n')
    text_paragraphs: List[str] = []
    image_paths: List[str] = []
    table_data_list: List[List[List[str]]] = []

    i = 0
    current_paragraph_lines: List[str] = []

    while i < len(lines):
        line = lines[i]

        # --- Identify images (Markdown format: ![alt](path)) ---
        # Use regex to match image links
        img_match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
        if img_match:
            img_path = img_match.group(1)
            if current_paragraph_lines:
                text_paragraphs.append('\n'.join(current_paragraph_lines))
                current_paragraph_lines = []
            # Simple validation of whether path exists or is network path (assume local path here)
            # You can add more complex validation as needed
            image_paths.append(img_path)
            i += 1
            continue

        # --- Identify tables (Markdown format) ---
        # Tables are usually separated by | and the second line is --- 
        # Check if current line and next line form table header
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|(\s*[-:]+\s*\|)+\s*$', lines[i+1]):
            if current_paragraph_lines:
                text_paragraphs.append('\n'.join(current_paragraph_lines))
                current_paragraph_lines = []

            table_data = []
            # Add header (current line)
            header_cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            table_data.append(header_cells)
            i += 2 # Skip header and separator line

            # Collect data rows
            while i < len(lines) and '|' in lines[i]:
                data_cells = [cell.strip() for cell in lines[i].split('|') if cell.strip()]
                # Ensure column count is consistent, pad with empty if insufficient
                while len(data_cells) < len(header_cells):
                    data_cells.append('')
                # If too many columns, truncate
                data_cells = data_cells[:len(header_cells)]
                table_data.append(data_cells)
                i += 1
            
            table_data_list.append(table_data)
            continue # Continue main loop, as i has been updated

        # --- Process text ---
        # Empty line indicates paragraph separator
        if not line.strip():
            if current_paragraph_lines:
                text_paragraphs.append('\n'.join(current_paragraph_lines))
                current_paragraph_lines = []
        else:
            current_paragraph_lines.append(line)
        i += 1

    # Process last paragraph
    if current_paragraph_lines:
        text_paragraphs.append('\n'.join(current_paragraph_lines))
    print(image_paths)
    return text_paragraphs, image_paths, table_data_list

# --- 2. DOCX export function ---
def export_docx(
    output_path: str,
    text_paragraphs: List[str],
    image_paths: List[str],
    table_data_list: List[List[List[str]]],
    image_width_inches: float = 4.0
):
    """Export parsed content as DOCX file."""
    try:
        doc = Document()
        # Ensure resource list length matches paragraph count
        # e.g.: Paragraph1 -> Image1 -> Table1 -> Paragraph2 -> ...
        # So we have N paragraphs, at most N-1 images and N-1 tables
        num_resources = len(text_paragraphs) - 1 if text_paragraphs else 0
        # If resources exceed expected, truncate; if less, pad with None
        image_paths_padded = (image_paths + [None] * max(0, num_resources - len(image_paths)))[:num_resources] if num_resources > 0 else []
        table_data_list_padded = (table_data_list + [None] * max(0, num_resources - len(table_data_list)))[:num_resources] if num_resources > 0 else []

        for i, paragraph_text in enumerate(text_paragraphs):
            if paragraph_text.strip():
                doc.add_paragraph(paragraph_text)
            
            # Insert corresponding resources after paragraph (except last paragraph)
            if i < num_resources:
                # Insert image
                img_path = image_paths_padded[i]
                if img_path:
                    img_path = os.path.abspath(img_path)  # Ensure path is absolute
                    print(img_path)
                    # Simple check if file exists (for local paths)
                    if not (img_path.startswith('http://') or img_path.startswith('https://')):
                        if not os.path.exists(img_path):
                            doc.add_paragraph(f"[Warning: Image not found '{img_path}']")
                        else:
                            try:
                                doc.add_paragraph() # Empty line
                                doc.add_picture(img_path, width=Inches(image_width_inches))
                            except Exception as e:
                                doc.add_paragraph(f"[Image insertion failed: {img_path} - {e}]")
                    else: # Network image, python-docx doesn't directly support, need to download first
                         doc.add_paragraph(f"[Network image: {img_path}]") # Placeholder

                # Insert table
                table_data = table_data_list_padded[i]
                if table_data:
                    try:
                        doc.add_paragraph() # Empty line
                        if table_data and len(table_data) > 0 and len(table_data[0]) > 0:
                            num_cols = len(table_data[0])
                            table = doc.add_table(rows=0, cols=num_cols)
                            table.style = 'Table Grid'
                            
                            # Add header and bold
                            hdr_cells = table.add_row().cells
                            for j, cell_value in enumerate(table_data[0]):
                                hdr_cells[j].text = str(cell_value)
                                # For header bold, can apply style or directly manipulate run
                                # Simple handling here, may need more complex style settings in practice
                            
                            # Add data rows
                            for row_data in table_data[1:]:
                                row_cells = table.add_row().cells
                                padded_row_data = (row_data + [''] * num_cols)[:num_cols]
                                for j, cell_value in enumerate(padded_row_data):
                                    row_cells[j].text = str(cell_value)
                        else:
                            doc.add_paragraph("[Table data is empty or invalid]")
                    except Exception as e:
                        doc.add_paragraph(f"[Table insertion failed: {e}]")

        doc.save(output_path)
        print(f"✅ DOCX file saved to: {output_path}")
        return True
    except Exception as e:
        print(f"❌ Error generating DOCX file: {e}")
        import traceback
        traceback.print_exc()
        return False




def convert_docx_to_pdf_with_libreoffice(
    input_docx_path: str,
    output_pdf_path: str,
    libreoffice_command: str = "soffice", # Try 'soffice' or 'libreoffice'
    overwrite: bool = True
) -> bool:
    """
    Use LibreOffice command-line tool to convert DOCX file to PDF file.

    This method meets the requirement mentioned in the knowledge base that "answer content supports text, well-formatted PDF/Word and other format files",
    and utilizes the LibreOffice tool installed on the system.

    Args:
        input_docx_path (str): Input DOCX file path.
        output_pdf_path (str): Output PDF file path.
        libreoffice_command (str): Command to call LibreOffice, default is 'soffice'.
                                   May need 'libreoffice' on some systems.
        overwrite (bool): If output file exists, LibreOffice will usually overwrite it.

    Returns:
        bool: Returns True if conversion successful, otherwise False.
    """
    # 1. Check if input file exists
    if not os.path.exists(input_docx_path):
        print(f"❌ Error: Input DOCX file '{input_docx_path}' does not exist.")
        return False

    # 2. Get and create output directory (if needed)
    output_dir = os.path.dirname(output_pdf_path)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"📁 Created output directory: {output_dir}")
        except OSError as e:
            print(f"❌ Error: Cannot create output directory '{output_dir}': {e}")
            return False

    # 3. Build LibreOffice command
    # --headless: Run in headless mode
    # --convert-to pdf: Specify conversion target format as PDF
    # --outdir: Specify output directory
    cmd = [
        libreoffice_command,
        "--headless",             # Headless mode
        "--convert-to", "pdf",    # Convert to PDF
        "--outdir", output_dir if output_dir else ".", # Output directory
        input_docx_path           # Input file
    ]

    print(f"🔄 Calling LibreOffice command: {' '.join(cmd)}")

    try:
        # 4. Execute command
        result = subprocess.run(
            cmd,
            check=True,           # If return code is non-zero, raise CalledProcessError
            capture_output=True,  # Capture stdout and stderr
            text=True,            # Decode output as string
            timeout=120           # Set timeout (seconds)
        )
        print(f"✅ LibreOffice conversion command executed successfully.")

        # 5. Verify output file exists
        if os.path.exists(output_pdf_path):
            print(f"✅ PDF file successfully generated: {output_pdf_path}")
            return True
        else:
            print(f"⚠️  LibreOffice command executed successfully, but PDF file not found at '{output_pdf_path}'.")
            # Can try listing output directory contents for debugging
            if output_dir:
                print(f"   Output directory '{output_dir}' contents:")
                try:
                    for f in os.listdir(output_dir):
                        print(f"     - {f}")
                except OSError:
                    pass
            return False

    except subprocess.CalledProcessError as e:
        print(f"❌ LibreOffice conversion failed (return code {e.returncode}): {e}")
        if e.stdout:
            print(f"   Stdout: {e.stdout}")
        if e.stderr:
            print(f"   Stderr: {e.stderr}")
        return False
    except subprocess.TimeoutExpired:
        print(f"❌ LibreOffice conversion timeout (exceeded 120 seconds).")
        return False
    except FileNotFoundError:
        print(f"❌ Command '{libreoffice_command}' not found. Please ensure LibreOffice is installed and command is in PATH.")
        print(f"   You may need to try using 'libreoffice' as the command.")
        return False
    except Exception as e:
        print(f"❌ Unknown error occurred when calling LibreOffice conversion: {e}")
        import traceback
        traceback.print_exc()
        return False

# --- Example usage ---
def export2pdf(input_docx, output_pdf):
    # --- Configuration ---
    # Please replace the paths below with your actual DOCX file paths
    # input_docx = "example_output.docx"  # Input DOCX file
    # output_pdf = "converted_output.pdf" # Output PDF file
    # If 'soffice' command doesn't work, try 'libreoffice'
    libreoffice_cmd = "soffice" # or "libreoffice"
    # --- End configuration ---

    if not os.path.exists(input_docx):
        print(f"⚠️  Example input file '{input_docx}' does not exist. Please generate a DOCX file first or modify the path.")
        # Can choose to exit or prompt user
        # sys.exit(1)

    print(f"📄 Preparing to convert '{input_docx}' to '{output_pdf}'...")
    success = convert_docx_to_pdf_with_libreoffice(
        input_docx_path=input_docx,
        output_pdf_path=output_pdf,
        libreoffice_command=libreoffice_cmd
    )

    if success:
        print(f"\n🎉 Conversion completed successfully!")
    else:
        print(f"\n💥 Conversion failed. Please check error messages.")
        # sys.exit(1) # Decide whether to exit based on needs

def md2docx(input_md, output_docx):
    print(f"📄 Preparing to convert '{input_md}' to '{output_docx}'...")
    # --- Read Markdown file ---
    try:
        with open(input_md, 'r', encoding='utf-8') as f:
            md_content = f.read()
        print(f"📄 Markdown file read: {input_md}")
    except Exception as e:
        print(f"❌ Error reading Markdown file: {e}")
        sys.exit(1)
    texts, images, tables = extract_content_from_md(md_content)

    # --- Generate file ---
    print("\n💾 Generating file...")
    success_docx = export_docx(output_docx, texts, images, tables)
    if success_docx:
        print(f"\n🎉 Conversion completed successfully!"
              "\n   - DOCX: {output_docx}")
    else:
        print(f"\n💥 Conversion failed. Please check error messages.")

def md2pdf(input_md, output_pdf):
    # Determine output file name
    base_name = os.path.splitext(input_md)[0]
    docx_path = f"{base_name}.docx"
   
    md2docx(input_md, docx_path)
    success_pdf = export2pdf(input_docx=docx_path, output_pdf=output_pdf)

    if success_pdf:
        print(f"\n🎉 Conversion completed successfully!"
              "\n   - PDF:  {output_pdf}")
    else:
        print(f"\n💥 Conversion failed. Please check error messages.")

def mdcontent2docx(md_content, output_docx):
    print(f"📄 Preparing to convert to '{output_docx}'...")
    texts, images, tables = extract_content_from_md(md_content)

    # --- Generate file ---
    print("\n💾 Generating file...")
    success_docx = export_docx(output_docx, texts, images, tables)
    if success_docx:
        print(f"\n🎉 Conversion completed successfully!"
              "\n   - DOCX: {output_docx}")
    else:
        print(f"\n💥 Conversion failed. Please check error messages.")

def mdcontent2pdf(mdcontent, output_pdf):
    # Determine output file name
    base_name = os.path.splitext(output_pdf)[0]
    docx_path = f"{base_name}.docx"
   
    mdcontent2docx(mdcontent, docx_path)
    success_pdf = export2pdf(input_docx=docx_path, output_pdf=output_pdf)

    if success_pdf:
        print(f"\n🎉 Conversion completed successfully!"
              "\n   - PDF:  {output_pdf}")
    else:
        print(f"\n💥 Conversion failed. Please check error messages.")

def mdcontent2md(mdcontent: str, md_path: str) -> bool:
    """
    Save the given Markdown string content to the specified file path.

    Args:
        mdcontent (str): String content containing Markdown syntax.
        md_path (str): Path to output Markdown file.

    Returns:
        bool: 如果文件保存成功返回 True，否则返回 False。
    """
    try:
        # --- 确保输出目录存在 ---
        output_dir = os.path.dirname(md_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"📁 已创建输出目录: {output_dir}")

        # --- 写入文件 ---
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(mdcontent)
        
        print(f"✅ Markdown 内容已成功保存至: {md_path}")
        return True

    except Exception as e:
        print(f"❌ 保存 Markdown 文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return False
    

# --- 4. 主程序 ---
def main():
    mdcontent = '''
            ### 语义图像构建方法

            #### 1. **固定部件构建流程**
            - **初始化张量**
                创建尺寸为 `width/8 × height/8`、通道数为 `item_class_num` 的灰度图像张量，初始值为 `0.01`（值域 `[0, 1]`），以节省内存。
            - **生成高斯椭圆**
                对每个固定部件的边界框，计算其中心点为均值 `μ`，半宽/半长为 `2σ`（X/Y方向独立），生成峰值为 `1.0` 的二维正态分布（公式见原文第2-4步）。
                *示例图：*
                ![图1：高斯椭圆分布](D:/adavance/tsy/rag4chat/output/test8/auto/images/258b7f7f1bcf4f9204aeb3191f283fc260e4d9d699d66ff3b1a854fc4c882aa2.jpg)        
            - **叠加语义分布**
                将各部件的高斯椭圆叠加至对应类别通道的图像中，重叠区域取相对高值（如图4所示）。

            #### 2. **旋转移动部件构建**
            - 文档提到其构建方法与固定部件**不一致**，但未提供具体步骤，需结合其他技术（如动态边界框调整）进一步研究。

            #### 3. **视觉分析参考**
            - 图像示例：
                ![图2：多通道语义分布](D:/adavance/tsy/rag4chat/output/test8/auto/images/258b7f7f1bcf4f9204aeb3191f283fc260e4d9d699d66ff3b1a854fc4c882aa2.jpg)      
                （注：此图为技术分析示例，包含红色标注的边界框和语义分布区域）'''
    mdcontent2pdf(mdcontent, "./example.pdf")
    # mdcontent2md(mdcontent, "./example0.md")
    # md2docx("./example0.md", "./example0.docx")
    # parser = argparse.ArgumentParser(description="将 Markdown 文件转换为 DOCX 和 PDF。")
    # parser.add_argument("input_md",default="./example.md", help="输入的 Markdown 文件路径")
    # parser.add_argument("-d", "--docx", help="输出的 DOCX 文件路径 (默认: input.md -> input.docx)")
    # parser.add_argument("-p", "--pdf", help="输出的 PDF 文件路径 (默认: input.md -> input.pdf)")
    
    # args = parser.parse_args()

    # input_md_path = args.input_md

    # if not os.path.exists(input_md_path):
    #     print(f"❌ 错误: 输入文件 '{input_md_path}' 不存在。")
    #     sys.exit(1)

    # # 确定输出文件名
    # base_name = os.path.splitext(input_md_path)[0]
    # output_docx_path = args.docx if args.docx else f"{base_name}.docx"
    # output_pdf_path = args.pdf if args.pdf else f"{base_name}.pdf"

    # # --- 读取 Markdown 文件 ---
    # try:
    #     with open(input_md_path, 'r', encoding='utf-8') as f:
    #         md_content = f.read()
    #     print(f"📄 已读取 Markdown 文件: {input_md_path}")
    # except Exception as e:
    #     print(f"❌ 读取 Markdown 文件时出错: {e}")
    #     sys.exit(1)

    # # --- 解析内容 ---
    # print("🔍 正在解析 Markdown 内容...")
    # texts, images, tables = extract_content_from_md(md_content)

    # print(f"  - 解析到文本段落数: {len(texts)}")
    # print(f"  - 解析到图片路径数: {len(images)}")
    # print(f"  - 解析到表格数量: {len(tables)}")

    # # --- 生成文件 ---
    # print("\n💾 正在生成文件...")
    # success_docx = export_docx(output_docx_path, texts, images, tables)
    # success_pdf = export2pdf(input_docx=output_docx_path, output_pdf=output_pdf_path)

    # if success_docx and success_pdf:
    #     print(f"\n🎉 所有文件已成功生成!")
    #     print(f"   - DOCX: {output_docx_path}")
    #     print(f"   - PDF:  {output_pdf_path}")
    # else:
    #     print(f"\n⚠️  部分文件生成失败。")
    #     sys.exit(1)

if __name__ == "__main__":
    main()